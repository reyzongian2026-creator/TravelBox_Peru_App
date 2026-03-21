from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "docs" / "TRAVELBOX_CASOS_DE_USO_Y_SECUENCIAS.docx"


def _add_title(document: Document, text: str) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)


def _add_meta(document: Document, text: str) -> None:
    p = document.add_paragraph(text)
    if p.runs:
        p.runs[0].italic = True


def _add_use_case(
    document: Document,
    code: str,
    title: str,
    actors: str,
    preconditions: list[str],
    happy: list[str],
    unhappy: list[str],
) -> None:
    document.add_heading(f"{code} - {title}", level=2)
    document.add_paragraph(f"Actores: {actors}")
    document.add_paragraph("Precondiciones:")
    for item in preconditions:
        document.add_paragraph(item, style="List Bullet")
    document.add_paragraph("Flujo Happy Path:")
    for item in happy:
        document.add_paragraph(item, style="List Number")
    document.add_paragraph("Flujo Unhappy / Excepciones:")
    for item in unhappy:
        document.add_paragraph(item, style="List Bullet")


def _add_sequence(document: Document, title: str, mermaid: str) -> None:
    document.add_heading(title, level=2)
    document.add_paragraph(
        "Diagrama de secuencia (Mermaid textual para documentación técnica):"
    )
    p = document.add_paragraph()
    run = p.add_run(mermaid)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def _add_screenshots(document: Document, screenshots: list[tuple[str, str]]) -> None:
    document.add_heading("Pantallazos Happy / Unhappy", level=2)
    for file_name, caption in screenshots:
        image_path = ROOT / file_name
        if not image_path.exists():
            continue
        document.add_paragraph(caption)
        try:
            document.add_picture(str(image_path), width=Inches(5.8))
        except Exception:
            document.add_paragraph(f"[No se pudo insertar imagen: {file_name}]")
        document.add_paragraph()


def build_document() -> None:
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _add_title(doc, "TravelBox - Casos de Uso, Secuencias y Validaciones")
    _add_meta(
        doc,
        f"Versión: {datetime.now().strftime('%Y-%m-%d %H:%M')} | "
        "Alcance: Cliente, Operador, Soporte, Courier y Admin",
    )

    doc.add_heading("1. Alcance homologado Front + Back", level=1)
    doc.add_paragraph(
        "Este documento consolida casos de uso funcionales y técnicos, "
        "incluyendo flujos happy/unhappy, validaciones clave, trazabilidad de "
        "incidencias por reserva, comportamiento en tiempo real y paginación latest-first."
    )

    doc.add_heading("2. Roles", level=1)
    for role in [
        "Cliente: crea y paga reservas, solicita recojo/delivery, reporta incidentes y da seguimiento.",
        "Operador: valida pagos caja, opera QR/PIN, registra equipaje y controla estado de reservas.",
        "Soporte: atiende y resuelve incidencias, mantiene trazabilidad operativa.",
        "Courier: toma servicios logísticos, actualiza estado y confirma entrega/recojo.",
        "Admin: gestiona almacenes/usuarios/roles, supervisa operaciones y analítica.",
    ]:
        doc.add_paragraph(role, style="List Bullet")

    doc.add_heading("3. Casos de Uso", level=1)
    _add_use_case(
        doc,
        "UC-CLI-001",
        "Cliente crea reserva sin recojo/delivery",
        "Cliente",
        [
            "Usuario autenticado con perfil completo.",
            "Sede con capacidad disponible.",
        ],
        [
            "Cliente selecciona sede y rango horario.",
            "Sistema valida capacidad y calcula precio.",
            "Sistema crea reserva en estado pendiente de pago.",
            "Cliente completa pago y reserva pasa a confirmada.",
            "Cliente visualiza QR y credenciales operativas.",
        ],
        [
            "Capacidad llena: sistema bloquea creación y muestra motivo.",
            "Pago rechazado: reserva queda pendiente y se notifica al cliente.",
        ],
    )
    _add_use_case(
        doc,
        "UC-CLI-002",
        "Cliente crea reserva con recojo",
        "Cliente, Operador, Courier",
        ["Cliente con ubicación válida y datos de contacto actualizados."],
        [
            "Cliente solicita recojo desde detalle de reserva.",
            "Sistema crea orden logística y notifica áreas operativas.",
            "Courier toma orden, actualiza estado y confirma recojo.",
            "Operador valida trazabilidad y flujo QR/PIN asociado.",
        ],
        [
            "Sin permisos de ubicación: sistema exige habilitar permiso o ubicación manual.",
            "No hay courier disponible: orden queda en cola y se mantiene trazabilidad.",
        ],
    )
    _add_use_case(
        doc,
        "UC-CLI-003",
        "Cliente crea reserva con delivery",
        "Cliente, Operador, Courier",
        ["Reserva en estado que permite entrega a domicilio."],
        [
            "Cliente registra dirección de entrega.",
            "Sistema genera orden de delivery y ETA estimado.",
            "Courier confirma tránsito y entrega.",
            "Operador cierra validaciones de identidad/equipaje/PIN.",
        ],
        [
            "PIN inválido: sistema bloquea cierre de entrega.",
            "Incidencia operativa: se abre ticket ligado a la reserva.",
        ],
    )
    _add_use_case(
        doc,
        "UC-CLI-004",
        "Cancelar reserva con reglas y motivo",
        "Cliente, Admin, Operador",
        ["Reserva existente y visible por el actor autorizado."],
        [
            "Usuario solicita cancelación.",
            "Sistema evalúa estado y tipo de pago.",
            "Si aplica, cancela o ejecuta reembolso + cancelación.",
            "Sistema registra motivo y notifica cambios.",
        ],
        [
            "Estado no cancelable: sistema informa motivo (completada/cancelada/expirada/en operación).",
            "Pago digital confirmado sin reembolso: sistema bloquea cancelación directa.",
        ],
    )
    _add_use_case(
        doc,
        "UC-INC-001",
        "Crear y dar seguimiento a incidencia por reserva",
        "Cliente, Operador, Soporte, Admin",
        ["Reserva identificada y visible por rol."],
        [
            "Actor abre incidencia vinculada a reserva.",
            "Sistema registra detalle + evidencia + trazabilidad.",
            "Sistema notifica audiencias por rol/sede.",
            "Soporte/Admin resuelve y sistema deja resolución auditable.",
        ],
        [
            "Reserva fuera de alcance del rol: acceso denegado.",
            "Incidencia ya resuelta: sistema evita doble cierre.",
        ],
    )
    _add_use_case(
        doc,
        "UC-OPS-001",
        "Operador valida pago en caja",
        "Operador, Cliente",
        ["Pago offline en estado pendiente."],
        [
            "Operador revisa intento de pago pendiente.",
            "Aprueba o rechaza con motivo.",
            "Sistema actualiza estado de pago y reserva.",
            "Cambio se refleja en tiempo real sin refresh manual.",
        ],
        [
            "Intento ya procesado: sistema bloquea doble acción.",
            "Método no offline: flujo de caja no aplica.",
        ],
    )
    _add_use_case(
        doc,
        "UC-OPS-002",
        "Flujo QR/PIN y registro de fotos de maletas",
        "Operador, Cliente, Courier",
        ["Reserva confirmada con acceso al módulo QR/PIN."],
        [
            "Operador escanea QR y valida reserva.",
            "Registra bag tag y fotos de equipaje.",
            "Genera/valida PIN según etapa.",
            "Sistema persiste trazabilidad y actualiza estado en vivo.",
        ],
        [
            "Faltan fotos requeridas: sistema bloquea avance.",
            "PIN inválido o no generado: sistema evita cierre.",
        ],
    )
    _add_use_case(
        doc,
        "UC-ADM-001",
        "Admin gestiona usuarios y sedes",
        "Admin",
        ["Rol ADMIN activo."],
        [
            "Admin accede a pestaña de operaciones de administración.",
            "Crea/edita usuarios y configuración de almacenes.",
            "Sistema valida reglas de negocio y actualiza catálogos.",
            "Cambios se reflejan en tiempo real en vistas dependientes.",
        ],
        [
            "Rol o datos inválidos: sistema rechaza operación con mensaje.",
            "Conflicto por correo existente: sistema informa causa.",
        ],
    )

    doc.add_heading("4. Secuencias técnicas", level=1)
    _add_sequence(
        doc,
        "SEQ-001 Reserva + Pago + Confirmación en tiempo real",
        """sequenceDiagram
participant C as Cliente
participant F as Frontend
participant B as Backend
participant N as NotificationService
participant O as Operador/Admin

C->>F: Crear reserva
F->>B: POST /reservations
B-->>F: Reserva PENDING_PAYMENT
C->>F: Confirmar pago
F->>B: POST /payments/confirm
B->>N: notifyPaymentConfirmed + evento realtime
N-->>F: SSE/stream notification
N-->>O: SSE/stream notification
F-->>C: Estado actualizado sin refresh""",
    )
    _add_sequence(
        doc,
        "SEQ-002 Incidencia por reserva con resolución",
        """sequenceDiagram
participant U as Usuario (Cliente/Operador)
participant F as Frontend
participant B as Backend
participant S as Soporte/Admin

U->>F: Crear incidencia
F->>B: POST /incidents
B-->>F: Ticket OPEN
B-->>S: Notificación por rol/sede
S->>F: Resolver incidencia
F->>B: PATCH /incidents/{id}/resolve
B-->>F: Ticket RESOLVED + trazabilidad""",
    )
    _add_sequence(
        doc,
        "SEQ-003 Paginación latest-first (5 en 5)",
        """sequenceDiagram
participant UI as UI
participant API as Backend API
UI->>API: GET /reservations/page?page=0&size=5
API-->>UI: items(desc), hasNext=true
UI->>API: GET next page
API-->>UI: siguientes 5 más recientes""",
    )

    doc.add_heading("5. Validaciones críticas", level=1)
    for item in [
        "No mezclar notificaciones entre usuarios/roles (limpieza por sesión + filtro de audiencia).",
        "Bloqueo explícito de cancelación según estado y tipo de pago.",
        "Ubicación obligatoria cuando cliente selecciona modo GPS actual.",
        "Validación de evidencia fotográfica en incidencias y flujo QR/PIN.",
        "Persistencia de idioma seleccionado entre pantallas y navegación.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    _add_screenshots(
        doc,
        [
            ("qa-login-screen.png", "Happy: autenticación y acceso inicial."),
            ("qa-post-login.png", "Happy: navegación posterior al login."),
            ("qa-before-submit.png", "Unhappy: validaciones previas al envío."),
            ("qa-valid-before-submit.png", "Happy: formulario válido antes de enviar."),
            ("qa-cash-page.png", "Operador/Admin: cola de pagos en caja."),
            ("qa-cash-after-approve.png", "Happy: pago en caja aprobado."),
            ("qa-ops-reservations.png", "Operador: reservas operativas."),
            ("qa-checkin-dialog.png", "Operación: check-in y evidencia."),
            ("qa-qr-presencial.png", "QR/PIN: flujo presencial."),
            ("qa-qr-presencial-after-pin.png", "QR/PIN: cierre con PIN."),
            ("qa-notifications-page.png", "Realtime: centro de notificaciones."),
        ],
    )

    doc.add_heading("6. Criterios de aceptación", level=1)
    for item in [
        "Cambios de estado visibles sin refresh manual en vistas principales.",
        "Paginación 5-en-5 en reservas, pagos e incidencias operativas.",
        "Textos sin hardcode en pantallas críticas; traducción consistente por locale.",
        "Trazabilidad completa de incidencias por reserva de apertura a resolución.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(str(OUT_PATH))
    print(f"Documento generado: {OUT_PATH}")


if __name__ == "__main__":
    build_document()
